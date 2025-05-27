#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_formatted_docx():
    # 创建新文档
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('告别AI代码生成的"胡言乱语"：MARIN框架如何让大模型不再编造API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加副标题
    subtitle = doc.add_paragraph('——基于层次依赖感知的API幻觉缓解技术研究')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True

    # 添加论文信息
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_run = info_para.add_run('论文地址：')
    info_run.bold = True
    info_para.add_run('https://arxiv.org/abs/2505.05057')

    # 添加摘要
    doc.add_heading('摘要', level=1)
    abstract = doc.add_paragraph(
        '大语言模型在代码生成过程中经常出现API幻觉问题，即生成不存在的API或错误使用现有API。'
        '本文介绍了MARIN框架，这是一种基于层次依赖感知的创新解决方案，通过静态分析项目依赖关系'
        '和约束解码技术，有效缓解了API幻觉现象。实验结果表明，MARIN在准确率方面提升了175%，'
        '幻觉率降低了80%，为AI辅助编程提供了重要的技术突破。'
    )

    # 添加关键词
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run('关键词：')
    keywords_run.bold = True
    keywords_para.add_run('大语言模型、API幻觉、代码生成、依赖分析、约束解码')

    doc.add_page_break()

    # 1. 引言
    doc.add_heading('1. 引言', level=1)
    doc.add_paragraph(
        '随着大语言模型（LLMs）在代码生成领域的广泛应用，开发者越来越依赖AI工具来辅助编程。'
        '然而，现有的代码生成模型主要基于预训练数据中的模式，往往忽略了实际项目中的具体依赖关系，'
        '导致生成的代码中出现API幻觉问题。'
    )

    doc.add_paragraph(
        'API幻觉主要表现为两种形式：一是调用项目中不存在的虚构API；二是错误使用已有API的调用方式。'
        '这些问题在实际开发中会导致代码无法正常运行，严重影响了AI辅助编程的实用性和可靠性。'
    )

    # 2. 问题分析
    doc.add_heading('2. 核心问题分析', level=1)

    doc.add_heading('2.1 API幻觉现象', level=2)
    doc.add_paragraph(
        '大语言模型在生成代码时面临的主要挑战包括：'
    )
    bullet_list = doc.add_paragraph('• 调用不存在的API（虚构API）', style='List Bullet')
    bullet_list = doc.add_paragraph('• 错误使用已有API（用法不当）', style='List Bullet')
    bullet_list = doc.add_paragraph('• 忽略项目特定的依赖关系', style='List Bullet')

    doc.add_heading('2.2 现有方法的局限性', level=2)
    doc.add_paragraph(
        '传统的RAG（检索增强生成）方法虽然能够提供相关代码片段，但存在以下不足：'
    )
    bullet_list = doc.add_paragraph('• 提供孤立的代码片段，缺乏结构性上下文', style='List Bullet')
    bullet_list = doc.add_paragraph('• 无法充分理解项目的依赖关系', style='List Bullet')
    bullet_list = doc.add_paragraph('• 维护检索语料库成本高，可扩展性差', style='List Bullet')

    doc.add_heading('2.3 实验发现', level=2)
    doc.add_paragraph(
        '研究发现，对于项目自定义的API方法，模型的幻觉率显著高于第三方标准API。'
        '这是因为第三方API在模型预训练阶段已经大量出现，而项目特定的API缺乏足够的训练数据支持。'
    )

    # 3. MARIN框架
    doc.add_heading('3. MARIN解决方案', level=1)

    doc.add_paragraph(
        'MARIN（Mitigating API Hallucination with Hierarchical Dependency Aware）是一个创新的框架，'
        '通过两个核心阶段有效缓解API幻觉问题：'
    )

    doc.add_heading('3.1 层次依赖挖掘', level=2)
    doc.add_paragraph(
        '该阶段使用静态分析技术解析项目的依赖关系，包括：'
    )

    doc.add_heading('项目描述', level=3)
    doc.add_paragraph('为LLM提供项目背景和目的的简要概述')

    doc.add_heading('全局依赖', level=3)
    doc.add_paragraph('提取相关文件的结构骨架，包括类定义、成员字段和函数签名')

    doc.add_heading('局部依赖', level=3)
    doc.add_paragraph('识别直接相关的函数调用和API引用')

    doc.add_heading('不完整函数', level=3)
    doc.add_paragraph('标记需要生成API调用的具体位置')

    doc.add_heading('3.2 依赖约束解码', level=2)
    doc.add_paragraph(
        '通过构建API名称前缀树和参数模式识别，在解码过程中动态约束生成的token，'
        '确保生成的API符合项目的实际依赖关系。这种方法能够：'
    )
    bullet_list = doc.add_paragraph('• 实时验证API的存在性', style='List Bullet')
    bullet_list = doc.add_paragraph('• 确保API调用的正确性', style='List Bullet')
    bullet_list = doc.add_paragraph('• 减少幻觉现象的发生', style='List Bullet')

    # 4. 实验验证
    doc.add_heading('4. 实验验证与结果', level=1)

    doc.add_heading('4.1 实验设计', level=2)
    doc.add_paragraph(
        '为了全面评估MARIN的效果，研究团队设计了comprehensive的实验方案：'
    )

    doc.add_heading('基准测试', level=3)
    doc.add_paragraph(
        '• APIHulBench：包含416个来自98个Java项目的高质量样本'
    )
    doc.add_paragraph(
        '• 工业场景基准：109个华为内部Java项目样本'
    )

    doc.add_heading('评估模型', level=3)
    doc.add_paragraph(
        '• CodeLlama（7B、13B、34B）'
    )
    doc.add_paragraph(
        '• DeepSeekCoder（1.3B、6.7B、33B）'
    )
    doc.add_paragraph(
        '• PanguCoder（11B、34B）'
    )

    doc.add_heading('4.2 关键指标', level=2)

    doc.add_heading('准确性指标', level=3)
    bullet_list = doc.add_paragraph('• Exact Match (EM)：完全匹配率', style='List Bullet')
    bullet_list = doc.add_paragraph('• Edit Similarity (ES)：编辑相似度', style='List Bullet')
    bullet_list = doc.add_paragraph('• Identifier Match (IM)：标识符匹配率', style='List Bullet')

    doc.add_heading('幻觉指标', level=3)
    bullet_list = doc.add_paragraph('• Micro Hallucination Number (MiHN)：幻觉元素平均数量', style='List Bullet')
    bullet_list = doc.add_paragraph('• Macro Hallucination Rate (MaHR)：幻觉API比例', style='List Bullet')

    doc.add_heading('4.3 实验结果', level=2)

    # 创建表格展示结果
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模型'
    hdr_cells[1].text = 'EM提升'
    hdr_cells[2].text = 'MaHR降低'
    hdr_cells[3].text = '效率提升'

    # 数据行
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'CodeLlama-7B'
    row1_cells[1].text = '175.39%'
    row1_cells[2].text = '80.15%'
    row1_cells[3].text = '0.022s开销'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'DeepSeekCoder'
    row2_cells[1].text = '107.32%'
    row2_cells[2].text = '67.31%'
    row2_cells[3].text = '高效扩展'

    row3_cells = table.rows[3].cells
    row3_cells[0].text = 'PanguCoder-11B'
    row3_cells[1].text = '71.79%'
    row3_cells[2].text = '58.21%'
    row3_cells[3].text = '0.030s开销'

    doc.add_paragraph()
    doc.add_paragraph(
        '实验结果表明，MARIN在所有测试模型上都取得了显著的性能提升，'
        '平均准确率提升超过100%，幻觉率降低超过60%。'
    )

    # 5. 技术优势
    doc.add_heading('5. 技术优势与创新点', level=1)

    doc.add_heading('5.1 核心创新', level=2)
    bullet_list = doc.add_paragraph('• 首次提出层次化依赖感知机制', style='List Bullet')
    bullet_list = doc.add_paragraph('• 创新的约束解码技术', style='List Bullet')
    bullet_list = doc.add_paragraph('• 无需维护额外检索语料库', style='List Bullet')

    doc.add_heading('5.2 实用价值', level=2)
    bullet_list = doc.add_paragraph('• 显著提升代码生成质量', style='List Bullet')
    bullet_list = doc.add_paragraph('• 降低开发调试成本', style='List Bullet')
    bullet_list = doc.add_paragraph('• 提高AI辅助编程的可靠性', style='List Bullet')

    # 6. 结论与展望
    doc.add_heading('6. 结论与未来展望', level=1)

    doc.add_paragraph(
        'MARIN框架通过创新的层次依赖感知和约束解码技术，有效解决了大语言模型代码生成中的API幻觉问题。'
        '该方法不仅在学术基准测试中表现优异，在工业场景中也展现出了强大的实用价值。'
    )

    doc.add_paragraph(
        '未来的研究方向包括：'
    )
    bullet_list = doc.add_paragraph('• 扩展到更多编程语言的支持', style='List Bullet')
    bullet_list = doc.add_paragraph('• 优化依赖分析的效率和准确性', style='List Bullet')
    bullet_list = doc.add_paragraph('• 集成到主流IDE和开发工具中', style='List Bullet')

    # 7. 参考文献
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph(
        '[1] Towards Mitigating API Hallucination in Code Generated by LLMs with Hierarchical Dependency Aware. '
        'arXiv:2505.05057, 2024.'
    )

    # 保存文档
    doc.save('MARIN框架_API幻觉缓解技术研究.docx')
    print('Word文档已成功创建：MARIN框架_API幻觉缓解技术研究.docx')

if __name__ == "__main__":
    create_formatted_docx() 